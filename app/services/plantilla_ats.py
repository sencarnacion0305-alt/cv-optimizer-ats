"""
Generador de Plantilla ATS.

Toma un CV (texto, normalmente extraído de un DOCX/PDF cuyo formato puede romper
el parsing de un ATS) y lo RE-EMITE en un DOCX limpio:

  - Una sola columna, lectura de arriba a abajo.
  - Secciones estándar (Resumen, Experiencia, Educación, Habilidades…).
  - Fechas con separador consistente.
  - Sin tablas, columnas, gráficos, barras de nivel ni iconos.
  - Fuente segura para ATS (Calibri).

No inventa contenido: preserva la información del CV original y solo la
reorganiza en una estructura que cualquier ATS puede leer sin errores.
"""

import io
import re
from typing import List

from app.services.parser_ats import (
    EMAIL_RE, PHONE_RE, LINKEDIN_RE, UBIC_RE, RANGO_RE, _detectar_nombre,
)


# ---------------------------------------------------------------------------
# Clasificación de secciones y nombres canónicos
# ---------------------------------------------------------------------------

_TIPOS = [
    ("resumen",         r"(professional\s*summary|summary|profile|perfil|resumen|"
                        r"objetivo|about\s*me|sobre\s*m[ií])"),
    ("experiencia",     r"(work\s*experience|professional\s*experience|experience|"
                        r"experiencia|employment|trayectoria|historial\s*laboral)"),
    ("educacion",       r"(education|educaci[oó]n|formaci[oó]n|academic)"),
    ("habilidades",     r"(technical\s*skills?|skills?|habilidades|competenc|"
                        r"conocimientos|technologies)"),
    ("certificaciones", r"(certificat|certificac|licenc|credential)"),
    ("idiomas",         r"(languages?|idiomas?)"),
    ("proyectos",       r"(projects?|proyectos?)"),
    ("logros",          r"(achievements?|logros|awards?|reconocimientos?)"),
    ("contacto",        r"(contact|contacto|datos\s*personales|informaci[oó]n\s*de\s*contacto)"),
]

_CANON = {
    "es": {
        "resumen": "RESUMEN PROFESIONAL", "experiencia": "EXPERIENCIA PROFESIONAL",
        "educacion": "EDUCACIÓN", "habilidades": "HABILIDADES",
        "certificaciones": "CERTIFICACIONES", "idiomas": "IDIOMAS",
        "proyectos": "PROYECTOS", "logros": "LOGROS DESTACADOS",
    },
    "en": {
        "resumen": "PROFESSIONAL SUMMARY", "experiencia": "PROFESSIONAL EXPERIENCE",
        "educacion": "EDUCATION", "habilidades": "SKILLS",
        "certificaciones": "CERTIFICATIONS", "idiomas": "LANGUAGES",
        "proyectos": "PROJECTS", "logros": "KEY ACHIEVEMENTS",
    },
}

_ES_HINTS = re.compile(
    r"\b(experiencia|educaci|habilidad|años|gesti[oó]n|desarrollo|empresa|"
    r"responsable|conocimientos|gracias|gestor)\b", re.I)
_EN_HINTS = re.compile(
    r"\b(experience|education|skills|years|management|development|company|"
    r"responsible|knowledge|achievements|engineer)\b", re.I)

_REEMPLAZOS = {
    "“": '"', "”": '"', "‘": "'", "’": "'",
    "–": "-", "—": "-", "•": "", "▪": "", "●": "",
    "■": "", "‣": "", "❖": "", "★": "", "✔": "",
    "→": "", "\t": " ",
}

# Verbos de accion frecuentes al inicio de un bullet de logro
_VERBO_BULLET = re.compile(
    r"^(led|managed|built|created|designed|developed|implemented|reduced|"
    r"increased|improved|launched|delivered|drove|optimized|achieved|resolved|"
    r"detected|mitigated|analyzed|coordinated|standardized|executed|investigated|"
    r"monitored|deployed|configured|performed|administered|supported|completed|"
    r"identified|correlated|represented|reviewed|hardened|mentored|automated|"
    r"lider|gestion|desarroll|implement|dise[nñ]|reduj|aument|mejor|coordin|"
    r"detect|analic|resolv|gestion|cre[ée]|logr[ée])", re.I)


def _idioma(texto: str) -> str:
    return "es" if len(_ES_HINTS.findall(texto)) >= len(_EN_HINTS.findall(texto)) else "en"


def _tipo_header(header: str) -> str:
    hl = header.lower()
    for tipo, patron in _TIPOS:
        if re.search(patron, hl):
            return tipo
    return ""


def _es_header_seccion(linea: str) -> str:
    """
    Devuelve el tipo de sección SOLO si la línea es un encabezado real conocido
    (corto, pocas palabras, y coincide con un tipo estándar). Evita confundir
    nombres, líneas de contacto o nombres de empresa con encabezados.
    """
    base = linea.strip().rstrip(":").strip()
    if not base or len(base) > 45 or len(base.split()) > 5:
        return ""
    if RANGO_RE.search(base):          # líneas con fechas no son encabezados
        return ""
    return _tipo_header(base)


def _limpiar(t: str) -> str:
    for a, b in _REEMPLAZOS.items():
        t = t.replace(a, b)
    t = re.sub(r"\s{2,}", " ", t)
    return t.strip(" \t-|·")


def _norm_fechas(linea: str) -> str:
    """Normaliza separadores de rango de fechas a guion simple con espacios."""
    return re.sub(r"\s*[–—]\s*", " - ", linea)


def _es_linea_contacto_simple(t: str) -> bool:
    return bool(
        EMAIL_RE.search(t) or LINKEDIN_RE.search(t) or PHONE_RE.search(t)
        or re.match(r"(location|ubicaci|phone|tel[eé]fono|email|correo|address|direcci)",
                    t, re.IGNORECASE))


def _parece_subtitulo(t: str) -> bool:
    """Línea de título de puesto / institución (no un bullet de logro)."""
    if RANGO_RE.search(t):
        return True
    if len(t) <= 60 and not t.rstrip().endswith(".") and not _VERBO_BULLET.search(t):
        return True
    return False


# ---------------------------------------------------------------------------
# Construcción del DOCX limpio
# ---------------------------------------------------------------------------

def _add_section_header(doc, titulo: str):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _add_paragraph(doc, texto: str):
    from docx.shared import Pt
    if not texto.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(_limpiar(texto))
    r.font.size = Pt(10.5)


def _add_subtitulo(doc, texto: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(_limpiar(_norm_fechas(texto)))
    r.bold = True
    r.font.size = Pt(10.5)


def _add_bullet(doc, texto: str):
    from docx.shared import Pt
    limpio = _limpiar(_norm_fechas(texto))
    if not limpio:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(limpio)
    r.font.size = Pt(10.5)


def _add_lineas(doc, lineas: List[str]):
    for linea in lineas:
        t = linea.strip()
        if not t:
            continue
        if _parece_subtitulo(t):
            _add_subtitulo(doc, t)
        else:
            _add_bullet(doc, t)


def _add_skills(doc, lineas: List[str]):
    """Emite las habilidades como bullets de una sola línea (sin tablas/barras)."""
    for linea in lineas:
        t = _limpiar(linea)
        if t:
            _add_bullet(doc, t)


def generar_plantilla_ats(texto_cv: str) -> bytes:
    """Reconstruye el CV en un DOCX limpio y compatible con ATS."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    lineas_todas = [l for l in texto_cv.splitlines()]
    lineas_strip = [l.strip() for l in lineas_todas if l.strip()]
    idioma = _idioma(texto_cv)
    canon = _CANON[idioma]

    # --- Datos de contacto ---
    nombre = _detectar_nombre(lineas_strip) or (lineas_strip[0] if lineas_strip else "")
    email  = EMAIL_RE.search(texto_cv)
    tel    = PHONE_RE.search(texto_cv)
    linked = LINKEDIN_RE.search(texto_cv)
    ubic   = UBIC_RE.search(texto_cv)

    contacto = []
    if email:  contacto.append(email.group(0))
    if tel:    contacto.append(tel.group(0).strip())
    if linked: contacto.append(linked.group(0))
    if ubic:   contacto.append(ubic.group(1).strip()[:60])

    # --- Documento ---
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Encabezado
    if nombre:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(_limpiar(nombre).upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x12, 0x2A, 0x4A)
    if contacto:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("  |  ".join(contacto))
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # --- Detectar secciones SOLO por encabezados reales conocidos ---
    secciones = []
    actual = None
    headline = ""          # subtítulo profesional bajo el nombre (si existe)
    nombre_lower = _limpiar(nombre).lower()
    for linea in lineas_strip:
        tipo = _es_header_seccion(linea)
        if tipo:
            actual = {"tipo": tipo, "header": linea, "lineas": []}
            secciones.append(actual)
        elif actual is not None:
            actual["lineas"].append(linea)
        else:
            # Bloque de encabezado (antes de la 1ª sección): capturar el headline
            ls = linea.strip()
            if (not headline and ls.lower() != nombre_lower
                    and not _es_linea_contacto_simple(ls) and len(ls) <= 90):
                headline = ls

    # Emitir el headline profesional bajo el contacto
    if headline:
        from docx.shared import Pt, RGBColor
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(_limpiar(headline))
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    contacto_vals = {v.lower() for v in contacto}

    for sec in secciones:
        if sec["tipo"] == "contacto":
            continue
        cuerpo = []
        for ls in sec["lineas"]:
            if ls.lower() == nombre_lower:
                continue
            if any(c in ls.lower() for c in contacto_vals if c):
                continue
            cuerpo.append(ls)
        if not cuerpo:
            continue

        titulo_sec = canon.get(sec["tipo"], _limpiar(sec["header"]).upper().rstrip(":"))
        _add_section_header(doc, titulo_sec)

        if sec["tipo"] == "resumen":
            _add_paragraph(doc, " ".join(cuerpo))
        elif sec["tipo"] == "habilidades":
            _add_skills(doc, cuerpo)
        else:
            _add_lineas(doc, cuerpo)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
