"""
Adapta un CV en formato DOCX a una vacante especifica.
- Reescribe el resumen incorporando todas las keywords de la vacante.
- Agrega las keywords faltantes a la seccion de habilidades.
- Todo el formato, diseno y estructura originales se preservan intactos.
"""

import io
import re
from typing import List

from app.services.adaptador import (
    _keywords_de, _analizar_cobertura, _normalizar, _kw_presente,
    _detectar_titulo_vacante, _titulo_en_cv, _PALABRAS_TITULO,
)
from app.services.ats_checker import (
    EMOJI_RE, ACRONIMOS_LARGOS, SIMBOLOS_RAROS, _empieza_con_verbo_accion,
)
from app.services.mejorador_bullets import sufijo_metrica, tiene_metrica


# ---------------------------------------------------------------------------
# Patrones de seccion
# ---------------------------------------------------------------------------

PATRON_RESUMEN = re.compile(
    r"(professional\s*summary|summary|resumen\s*profesional|"
    r"perfil\s*profesional|career\s*objective|objetivo|about\s*me|profile)",
    re.IGNORECASE,
)

PATRON_HABILIDADES = re.compile(
    r"(technical\s*skills?|skills?|habilidades?|competencias?|"
    r"core\s*competencies|technologies?|tools?\s*&?\s*tech|conocimientos?)",
    re.IGNORECASE,
)

PATRON_SIGUIENTE_SECCION = re.compile(
    r"(work\s*experience|experiencia|education|educacion|certif|"
    r"projects?|proyectos?|references?|awards?|languages?|idiomas?|contact|publications?)",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Formato de keywords (mayusculas correctas para legibilidad)
# ---------------------------------------------------------------------------

_ACRONIMOS = {
    "siem", "soar", "edr", "xdr", "ids", "ips", "waf", "dlp", "soc", "oscp",
    "cissp", "cism", "ceh", "giac", "aws", "gcp", "iam", "pam", "sql", "nosql",
    "api", "rest", "dns", "vpn", "tls", "ssl", "sox", "gdpr", "pci", "hipaa",
    "nist", "mitre", "owasp", "ldap", "ssh", "ftp", "smtp", "http", "https",
    "tcp", "udp", "jwt", "ueba", "mfa", "sso", "ics", "csirt", "ioc", "ttp",
    "iso", "ai", "ml", "cis", "ngfw", "dast", "sast",
    "rgpd", "ens", "lopd", "nis2", "dora", "enisa",
}

_CASING_ESPECIAL = {
    "logscale": "LogScale", "logrhythm": "LogRhythm", "fortianalyzer": "FortiAnalyzer",
    "crowdstrike": "CrowdStrike", "sentinelone": "SentinelOne", "powershell": "PowerShell",
    "javascript": "JavaScript", "typescript": "TypeScript", "postgresql": "PostgreSQL",
    "mongodb": "MongoDB", "mysql": "MySQL", "github": "GitHub", "gitlab": "GitLab",
    "devops": "DevOps", "devsecops": "DevSecOps", "mlops": "MLOps", "powerbi": "PowerBI",
    "fastapi": "FastAPI", "nestjs": "NestJS", "opensearch": "OpenSearch",
    "qradar": "QRadar",
}


# Conectores que van en minuscula dentro de una frase (espanol e ingles)
_CONECTORES = {"de", "del", "a", "la", "el", "los", "las", "y", "en", "para",
               "con", "of", "and", "the", "to", "in", "for"}


def _formato_keyword(kw: str) -> str:
    """Devuelve la keyword con mayusculas profesionales (EDR, OSCP, Splunk, LogScale)."""
    palabras = kw.split()
    out = []
    for i, w in enumerate(palabras):
        wl = w.lower()
        if wl in _CASING_ESPECIAL:
            out.append(_CASING_ESPECIAL[wl])
        elif wl in _ACRONIMOS:
            out.append(wl.upper())
        elif i > 0 and wl in _CONECTORES:
            out.append(wl)
        else:
            out.append(w.capitalize())
    return " ".join(out)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _es_heading(parrafo) -> bool:
    estilo = (parrafo.style.name if parrafo.style else "").lower()
    if "heading" in estilo or "title" in estilo:
        return True
    texto = parrafo.text.strip()
    if not texto:
        return False
    if texto.isupper() and 3 <= len(texto) <= 60:
        return True
    if len(texto) < 60 and parrafo.runs:
        if all(run.bold for run in parrafo.runs if run.text.strip()):
            return True
    return False


def _reemplazar_texto_parrafo(parrafo, nuevo_texto: str):
    """Reemplaza el texto preservando el formato del primer run."""
    if not parrafo.runs:
        parrafo.text = nuevo_texto
        return
    r0 = parrafo.runs[0]
    fuente  = r0.font.name
    tamano  = r0.font.size
    negrita = r0.bold
    cursiva = r0.italic
    try:
        color = r0.font.color.rgb if r0.font.color.type else None
    except Exception:
        color = None

    for run in parrafo.runs:
        run.text = ""
    parrafo.runs[0].text = nuevo_texto

    r = parrafo.runs[0]
    if fuente:  r.font.name  = fuente
    if tamano:  r.font.size  = tamano
    if negrita is not None: r.bold   = negrita
    if cursiva is not None: r.italic = cursiva
    if color:   r.font.color.rgb = color


def _agregar_parrafo_con_estilo(doc, referencia, texto: str):
    """Agrega un parrafo nuevo copiando el estilo de un parrafo de referencia."""

    nuevo = doc.add_paragraph()
    # Copiar estilo del parrafo de referencia
    if referencia.style:
        nuevo.style = referencia.style
    if referencia.runs:
        run = nuevo.add_run(texto)
        r0  = referencia.runs[0]
        if r0.font.name:  run.font.name  = r0.font.name
        if r0.font.size:  run.font.size  = r0.font.size
        run.bold   = r0.bold
        run.italic = r0.italic
        try:
            if r0.font.color.type:
                run.font.color.rgb = r0.font.color.rgb
        except Exception:
            pass
    else:
        nuevo.add_run(texto)

    # Mover el parrafo justo despues del parrafo de referencia
    ref_elem  = referencia._element
    ref_parent = ref_elem.getparent()
    ref_idx   = list(ref_parent).index(ref_elem)
    ref_parent.remove(nuevo._element)
    ref_parent.insert(ref_idx + 1, nuevo._element)
    return nuevo


# ---------------------------------------------------------------------------
# Construccion del resumen con TODAS las keywords
# ---------------------------------------------------------------------------

def _construir_resumen_completo(
    texto_original: str,
    cubiertas: List[str],
    sugeridas: List[str],
) -> str:
    """
    Reescribe el resumen incorporando todas las keywords de la vacante.
    Usa el texto original como base y agrega lo que falta.
    """
    base = texto_original.strip()
    if not base:
        base = "Cybersecurity professional with experience in threat detection and incident response."

    texto_n = _normalizar(base)

    # Priorizar las keywords que faltan (sugeridas) antes que las ya cubiertas
    todas_kw   = sugeridas + cubiertas
    faltantes  = [kw for kw in todas_kw if not _kw_presente(texto_n, kw)]

    if not faltantes:
        return base

    # Dividir en grupos para insertarlos naturalmente (con formato profesional)
    grupo1 = [_formato_keyword(k) for k in faltantes[:4]]   # tecnologias prioritarias
    grupo2 = [_formato_keyword(k) for k in faltantes[4:8]]  # otras keywords

    adaptado = base.rstrip(".")

    if grupo1:
        adaptado += f", with hands-on expertise in {', '.join(grupo1)}"
    if grupo2:
        adaptado += f" and proficiency in {', '.join(grupo2)}"

    adaptado += "."

    return adaptado


# ---------------------------------------------------------------------------
# Funcion principal
# ---------------------------------------------------------------------------

def _todos_los_parrafos_docx(doc):
    """
    Devuelve TODOS los parrafos del documento en orden:
    cuerpo principal + todas las celdas de todas las tablas.
    Cada elemento es el objeto Paragraph de python-docx.
    """
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxParagraph

    parrafos = []
    body = doc.element.body

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            parrafos.append(DocxParagraph(child, doc))
        elif tag == "tbl":
            for row in child.iter(qn("w:tr")):
                for cell in row.iter(qn("w:tc")):
                    for p_elem in cell.iter(qn("w:p")):
                        parrafos.append(DocxParagraph(p_elem, doc))

    return parrafos


# ---------------------------------------------------------------------------
# Auto-correccion de legibilidad ATS
# ---------------------------------------------------------------------------

_REEMPLAZOS_CHARS = {
    "“": '"', "”": '"', "‘": "'", "’": "'",
    "‚": "'", "„": '"', "«": '"', "»": '"',
    "–": "-", "—": "-", "―": "-", "−": "-",
}
_SIMBOLOS_INICIO = re.compile(r"^[" + re.escape(SIMBOLOS_RAROS) + r"]\s*")
_SIMBOLOS_TODOS  = re.compile(r"[" + re.escape(SIMBOLOS_RAROS) + r"]")


def _limpiar_caracteres_ats(t: str) -> str:
    for a, b in _REEMPLAZOS_CHARS.items():
        t = t.replace(a, b)
    t = _SIMBOLOS_INICIO.sub("", t)   # vineta decorativa al inicio
    t = _SIMBOLOS_TODOS.sub("", t)    # simbolos decorativos sueltos
    t = EMOJI_RE.sub("", t)           # emojis
    return t


def _nodos_texto_run(run):
    """Nodos w:t del run. Editarlos directamente preserva tabs y saltos
    (asignar run.text destruye los elementos <w:tab/> del run)."""
    from docx.oxml.ns import qn
    return run._r.findall(qn("w:t"))


def _normalizar_caracteres_doc(doc) -> int:
    """Reemplaza comillas/guiones tipograficos, simbolos decorativos y emojis."""
    cambios = 0
    for p in _todos_los_parrafos_docx(doc):
        for r in p.runs:
            for nodo in _nodos_texto_run(r):
                if nodo.text:
                    nuevo = _limpiar_caracteres_ats(nodo.text)
                    if nuevo != nodo.text:
                        nodo.text = nuevo
                        cambios += 1
    return cambios


def _expandir_acronimos_doc(doc, texto_completo: str) -> int:
    """Expande la primera aparicion de cada acronimo conocido: 'Forma Larga (ACR)'."""
    texto_low = texto_completo.lower()
    pendientes = {
        acr: largo for acr, largo in ACRONIMOS_LARGOS.items()
        if re.search(rf"\b{re.escape(acr)}\b", texto_low) and largo.lower() not in texto_low
    }
    if not pendientes:
        return 0

    ya = set()
    cambios = 0
    for p in _todos_los_parrafos_docx(doc):
        for r in p.runs:
            for nodo in _nodos_texto_run(r):
                if not nodo.text:
                    continue
                for acr, largo in list(pendientes.items()):
                    if acr in ya:
                        continue
                    patron = re.compile(rf"\b{re.escape(acr)}\b", re.IGNORECASE)
                    m = patron.search(nodo.text)
                    if m:
                        orig = m.group(0)
                        nodo.text = (nodo.text[:m.start()]
                                     + f"{largo} ({orig.upper()})"
                                     + nodo.text[m.end():])
                        ya.add(acr)
                        cambios += 1
    return cambios


def _append_a_parrafo(parrafo, texto_extra: str):
    """Agrega texto al final de un parrafo preservando el formato del ultimo run."""
    runs = parrafo.runs
    if not runs:
        parrafo.add_run(texto_extra)
        return
    ultimo = runs[-1]
    # Quitar el punto final del ultimo run (lo re-agregamos al final del extra)
    if ultimo.text.rstrip().endswith("."):
        ultimo.text = ultimo.text.rstrip()[:-1]
    nuevo = parrafo.add_run(texto_extra)
    try:
        if ultimo.font.name:  nuevo.font.name = ultimo.font.name
        if ultimo.font.size:  nuevo.font.size = ultimo.font.size
        nuevo.bold   = ultimo.bold
        nuevo.italic = ultimo.italic
        if ultimo.font.color and ultimo.font.color.type:
            nuevo.font.color.rgb = ultimo.font.color.rgb
    except Exception:
        pass


PATRON_EXPERIENCIA = re.compile(
    r"(work\s*experience|professional\s*experience|employment|"
    r"experiencia(\s+(laboral|profesional))?|trayectoria|historial\s*laboral)",
    re.IGNORECASE)


def _es_heading_seccion(parrafo) -> bool:
    """True si el parrafo es un encabezado de seccion (no un titulo de puesto)."""
    t = parrafo.text.strip()
    if not t or len(t) >= 45:
        return False
    return (_es_heading(parrafo)
            and (PATRON_EXPERIENCIA.search(t) or PATRON_SIGUIENTE_SECCION.search(t)
                 or PATRON_RESUMEN.search(t) or PATRON_HABILIDADES.search(t)))


def _agregar_metricas_doc(doc) -> int:
    """
    Agrega una metrica de impacto a los bullets de logro sin numeros,
    SOLO dentro de la seccion de experiencia laboral (evita certs, educacion, skills).
    """
    contadores: dict = {}
    cambios = 0
    en_experiencia = False

    for p in _todos_los_parrafos_docx(doc):
        t = p.text.strip()
        if not t:
            continue

        # Cambios de seccion
        if _es_heading_seccion(p):
            en_experiencia = bool(PATRON_EXPERIENCIA.search(t))
            continue

        if not en_experiencia:
            continue
        if not (40 <= len(t) <= 350):
            continue
        if _es_heading(p) or not _empieza_con_verbo_accion(t) or tiene_metrica(t):
            continue

        _append_a_parrafo(p, ", " + sufijo_metrica(t, contadores) + ".")
        cambios += 1
    return cambios


def _insertar_titulo_objetivo(doc, titulo: str) -> bool:
    """Agrega el titulo del puesto al titular profesional si no esta presente."""
    if not titulo:
        return False
    parrafos = [p for p in _todos_los_parrafos_docx(doc) if p.text.strip()]
    if not parrafos:
        return False
    titulo_n = _normalizar(titulo)

    # Si ya aparece en las primeras lineas, no hacer nada
    for p in parrafos[:5]:
        if titulo_n in _normalizar(p.text):
            return False

    # Headline = primera linea (tras el nombre) con '|' o palabra de cargo, sin contacto
    for p in parrafos[1:4]:
        t = p.text.strip()
        if ("|" in t or _PALABRAS_TITULO.search(t)) and "@" not in t and not re.search(r"\d{3}", t):
            _reemplazar_texto_parrafo(p, t.rstrip(" |") + " | " + titulo)
            return True
    return False


def adaptar_cv_docx(docx_bytes: bytes, vacante_texto: str) -> bytes:
    """
    Modifica el DOCX original:
    1. Reescribe el resumen con todas las keywords de la vacante.
    2. Agrega las keywords faltantes a la seccion de habilidades.
    El formato, diseno y estructura del documento se preservan intactos.
    Soporta CVs con layout en tablas.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    # Usar extraccion completa (parrafos + tablas)
    todos = _todos_los_parrafos_docx(doc)
    texto_cv_completo = "\n".join(p.text for p in todos if p.text.strip())

    kw_vacante           = _keywords_de(vacante_texto)[:30]
    cubiertas, sugeridas = _analizar_cobertura(kw_vacante, texto_cv_completo)

    # ── Paso 0: Job title match — insertar el cargo objetivo si falta ──
    titulo_vacante = _detectar_titulo_vacante(vacante_texto)
    if titulo_vacante and not _titulo_en_cv(titulo_vacante, texto_cv_completo):
        _insertar_titulo_objetivo(doc, titulo_vacante)
        todos = _todos_los_parrafos_docx(doc)  # refrescar tras el cambio

    # ── Paso 1: Adaptar el resumen ────────────────────────────────
    en_resumen       = False
    parrafos_resumen = []

    for parrafo in todos:
        texto = parrafo.text.strip()

        if _es_heading(parrafo) or (len(texto) < 60 and (
                PATRON_RESUMEN.match(texto) or PATRON_SIGUIENTE_SECCION.match(texto) or
                PATRON_HABILIDADES.match(texto))):
            if en_resumen:
                break
            if PATRON_RESUMEN.search(texto):
                en_resumen = True
            continue

        if en_resumen and texto:
            parrafos_resumen.append(parrafo)

    # Fallback: primer parrafo largo despues de datos de contacto
    if not parrafos_resumen:
        contacto_visto = False
        for parrafo in todos:
            texto = parrafo.text.strip()
            if not texto:
                continue
            if re.search(r"@|linkedin|\+\d|\d{3}.*\d{4}|location|phone|email", texto, re.I):
                contacto_visto = True
                continue
            if contacto_visto and len(texto.split()) > 15 and not _es_heading(parrafo):
                parrafos_resumen.append(parrafo)
                break

    if parrafos_resumen:
        texto_original = " ".join(p.text.strip() for p in parrafos_resumen if p.text.strip())
        texto_adaptado = _construir_resumen_completo(texto_original, cubiertas, sugeridas)
        _reemplazar_texto_parrafo(parrafos_resumen[0], texto_adaptado)
        for p_extra in parrafos_resumen[1:]:
            _reemplazar_texto_parrafo(p_extra, "")

    # ── Paso 2: Agregar keywords faltantes a habilidades ─────────
    todos_actualizado = _todos_los_parrafos_docx(doc)
    texto_cv_actual   = "\n".join(p.text for p in todos_actualizado if p.text.strip())
    texto_n_actual    = _normalizar(texto_cv_actual)
    kw_faltantes      = [kw for kw in (cubiertas + sugeridas)
                         if not _kw_presente(texto_n_actual, kw)]

    if kw_faltantes:
        en_habilidades        = False
        ultimo_parrafo_skills = None

        for parrafo in todos_actualizado:
            texto = parrafo.text.strip()

            if _es_heading(parrafo) or (len(texto) < 60 and (
                    PATRON_HABILIDADES.match(texto) or PATRON_SIGUIENTE_SECCION.match(texto))):
                if en_habilidades:
                    break
                if PATRON_HABILIDADES.search(texto):
                    en_habilidades = True
                continue

            if en_habilidades and texto:
                ultimo_parrafo_skills = parrafo

        kw_faltantes_fmt = [_formato_keyword(k) for k in kw_faltantes]
        if ultimo_parrafo_skills is not None:
            # Agregar keywords faltantes al final de la seccion de habilidades
            texto_nuevo = " | ".join(kw_faltantes_fmt)
            _agregar_parrafo_con_estilo(doc, ultimo_parrafo_skills, texto_nuevo)
        else:
            # Si no se encontro seccion de habilidades, agregar al final del documento
            p = doc.add_paragraph()
            p.add_run("Additional Skills: " + " | ".join(kw_faltantes_fmt))

    # ── Paso 3: Agregar metricas de impacto a bullets sin numeros ──
    _agregar_metricas_doc(doc)

    # ── Paso 4: Auto-corregir legibilidad ATS ────────────────────
    # Normalizar caracteres problematicos (comillas/guiones tipograficos, simbolos, emojis)
    _normalizar_caracteres_doc(doc)
    # Expandir acronimos conocidos la primera vez (SIEM -> Security Information... (SIEM))
    texto_final = "\n".join(p.text for p in _todos_los_parrafos_docx(doc) if p.text.strip())
    _expandir_acronimos_doc(doc, texto_final)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
