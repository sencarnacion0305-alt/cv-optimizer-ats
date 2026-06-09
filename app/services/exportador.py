"""
Exportador de CV adaptado a PDF y DOCX con formato profesional.
"""

import io
import re

from app.services.generador_cv import generar_cv_adaptado


def _preparar_datos(cv_texto: str, vacante_texto: str) -> dict:
    return generar_cv_adaptado(cv_texto, vacante_texto)


def _secciones_del_texto(texto: str):
    bloques = []
    lineas = texto.splitlines()

    SECCIONES_CONOCIDAS = {
        "RESUMEN PROFESIONAL", "EXPERIENCIA PROFESIONAL",
        "HABILIDADES TECNICAS", "HABILIDADES TECNICAS",
        "EDUCACION", "CERTIFICACIONES", "RESUMEN",
        "EXPERIENCIA", "HABILIDADES", "SKILLS",
        "PROFESSIONAL SUMMARY", "WORK EXPERIENCE",
    }

    i = 0
    encabezado_terminado = False
    lineas_encabezado = 0

    while i < len(lineas):
        stripped = lineas[i].strip()
        i += 1

        if not stripped:
            continue

        if re.match(r"^-{10,}$", stripped):
            continue

        if stripped.upper() in SECCIONES_CONOCIDAS:
            encabezado_terminado = True
            bloques.append({"tipo": "seccion", "texto": stripped.upper()})
            continue

        if stripped.startswith("•") or stripped.startswith(">") and len(stripped) > 2:
            encabezado_terminado = True
            texto_bullet = re.sub(r"^[•>\*\-]\s*", "", stripped)
            bloques.append({"tipo": "bullet", "texto": texto_bullet})
            continue

        if stripped.startswith("[") and stripped.endswith("]"):
            bloques.append({"tipo": "sugerencia", "texto": stripped[1:-1].strip()})
            continue

        if not encabezado_terminado:
            if lineas_encabezado == 0:
                bloques.append({"tipo": "nombre", "texto": stripped})
            elif lineas_encabezado == 1:
                bloques.append({"tipo": "subtitulo", "texto": stripped})
            else:
                bloques.append({"tipo": "contacto", "texto": stripped})
            lineas_encabezado += 1
            continue

        bloques.append({"tipo": "texto", "texto": stripped})

    return bloques


def _safe(t: str) -> str:
    # Reemplaza caracteres Unicode no soportados por Helvetica/latin-1
    t = t.replace("–", "-")   # en-dash
    t = t.replace("—", "-")   # em-dash
    t = t.replace("‘", "'")   # comilla simple izq
    t = t.replace("’", "'")   # comilla simple der
    t = t.replace("“", '"')   # comilla doble izq
    t = t.replace("”", '"')   # comilla doble der
    t = t.replace("•", ">")   # bullet
    t = t.replace("\U0001f4a1", ">>")  # emoji bombilla
    return t.encode("latin-1", errors="replace").decode("latin-1")


# ---------------------------------------------------------------------------
# Exportar a DOCX
# ---------------------------------------------------------------------------

def exportar_docx(cv_texto: str, vacante_texto: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    datos   = _preparar_datos(cv_texto, vacante_texto)
    bloques = _secciones_del_texto(datos["texto_completo"])
    doc     = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def _linea_decorativa(doc):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "6C63FF")
        pBdr.append(bot)
        pPr.append(pBdr)

    for bloque in bloques:
        tipo  = bloque["tipo"]
        texto = bloque["texto"]

        if tipo == "nombre":
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(texto)
            run.bold           = True
            run.font.size      = Pt(22)
            run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

        elif tipo == "subtitulo":
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(texto)
            run.font.size      = Pt(12)
            run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)

        elif tipo == "contacto":
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(texto)
            run.font.size      = Pt(9.5)
            run.font.color.rgb = RGBColor(0x88, 0x92, 0xA4)

        elif tipo == "seccion":
            doc.add_paragraph()
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(texto)
            run.bold           = True
            run.font.size      = Pt(11)
            run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
            _linea_decorativa(doc)

        elif tipo == "bullet":
            p   = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(texto)
            run.font.size = Pt(10.5)

        elif tipo == "sugerencia":
            p   = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(">> " + texto)
            run.italic         = True
            run.font.size      = Pt(9.5)
            run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)

        else:
            p   = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(texto)
            run.font.size = Pt(10.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Exportar a PDF
# ---------------------------------------------------------------------------

def exportar_pdf(cv_texto: str, vacante_texto: str) -> bytes:
    from fpdf import FPDF

    datos   = _preparar_datos(cv_texto, vacante_texto)
    bloques = _secciones_del_texto(datos["texto_completo"])

    class CV_PDF(FPDF):
        def header(self): pass
        def footer(self): pass

    pdf = CV_PDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 15, 20)

    # Colores
    PURPLE       = (108, 99,  255)
    PURPLE_LIGHT = (100, 80,  220)   # mas oscuro para leer sobre blanco
    DARK_TEXT    = (30,  30,  35)    # texto principal oscuro sobre fondo blanco
    MUTED        = (100, 110, 125)
    BG_DARK      = (15,  17,  23)
    ORANGE       = (200, 120,  10)   # naranja mas oscuro para leer sobre blanco

    # Agrupar bloques de texto consecutivos en parrafos
    bloques_agrupados = []
    i = 0
    while i < len(bloques):
        b = bloques[i]
        if b["tipo"] == "texto":
            # Recolectar todas las lineas de texto seguidas
            partes = [b["texto"]]
            j = i + 1
            while j < len(bloques) and bloques[j]["tipo"] == "texto":
                partes.append(bloques[j]["texto"])
                j += 1
            bloques_agrupados.append({"tipo": "parrafo", "texto": " ".join(partes)})
            i = j
        else:
            bloques_agrupados.append(b)
            i += 1

    for bloque in bloques_agrupados:
        tipo  = bloque["tipo"]
        texto = _safe(bloque["texto"])

        if tipo == "nombre":
            pdf.set_fill_color(*BG_DARK)
            pdf.set_text_color(*PURPLE)
            pdf.set_font("Helvetica", "B", 20)
            pdf.cell(0, 13, texto, new_x="LMARGIN", new_y="NEXT", fill=True)

        elif tipo == "subtitulo":
            pdf.set_fill_color(*BG_DARK)
            pdf.set_text_color(167, 139, 250)
            pdf.set_font("Helvetica", "", 11)
            pdf.cell(0, 8, texto, new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(4)

        elif tipo == "contacto":
            pdf.set_text_color(*MUTED)
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 5, texto, new_x="LMARGIN", new_y="NEXT")

        elif tipo == "seccion":
            pdf.ln(5)
            pdf.set_text_color(*PURPLE_LIGHT)
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 7, texto, new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(*PURPLE)
            pdf.set_line_width(0.4)
            y = pdf.get_y()
            pdf.line(20, y, 190, y)
            pdf.ln(3)

        elif tipo == "bullet":
            pdf.set_text_color(*DARK_TEXT)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_x(26)
            pdf.multi_cell(170, 5.5, "- " + texto, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)

        elif tipo == "parrafo":
            pdf.set_text_color(*DARK_TEXT)
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5.5, texto, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

        elif tipo == "sugerencia":
            pdf.ln(2)
            pdf.set_text_color(*ORANGE)
            pdf.set_font("Helvetica", "I", 9)
            pdf.multi_cell(0, 5, texto, new_x="LMARGIN", new_y="NEXT")

    return bytes(pdf.output())
