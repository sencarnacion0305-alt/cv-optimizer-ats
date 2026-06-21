"""
Analizador de compatibilidad ATS de un CV (sin necesidad de vacante).

Inspecciona el archivo a nivel de formato (XML del DOCX) y de contenido para
detectar los problemas que hacen que un CV sea rechazado por los sistemas de
seguimiento de candidatos (ATS) antes de que un humano lo lea.

Categorias evaluadas:
  - Formato:    tablas, columnas, imagenes, text boxes, contacto en header/footer.
  - Estructura: presencia de las secciones clave (Contact, Summary, Experience,
                Education, Skills).
  - Contacto:   email, telefono y LinkedIn presentes y en el cuerpo.
  - Contenido:  verbos de accion debiles, bullets sin metricas, longitud.
"""

import io
import re
from typing import Dict, List, Tuple

from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Patrones reutilizables
# ---------------------------------------------------------------------------

EMAIL_RE    = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE    = re.compile(r"(?:\+?\d[\d\s().\-]{7,}\d)")
LINKEDIN_RE = re.compile(r"linkedin\.com/in/", re.IGNORECASE)

SECCIONES = {
    "Contacto": re.compile(
        r"(contact|datos\s*de\s*contacto|informaci[oó]n\s*de\s*contacto)", re.I),
    "Resumen / Perfil": re.compile(
        r"(professional\s*summary|summary|profile|resumen|perfil|objetivo|about)", re.I),
    "Experiencia": re.compile(
        r"(work\s*experience|experience|experiencia|employment|trayectoria)", re.I),
    "Educación": re.compile(
        r"(education|educaci[oó]n|formaci[oó]n\s*acad)", re.I),
    "Habilidades": re.compile(
        r"(technical\s*skills?|skills?|habilidades|competenc|conocimientos)", re.I),
}

# Frases / verbos debiles al inicio de un bullet (evitar en CV de impacto)
VERBOS_DEBILES = [
    "responsible for", "duties included", "tasked with", "worked on",
    "helped", "assisted with", "assisted in", "participated in",
    "involved in", "in charge of", "responsible of", "responsable de",
    "encargado de", "encargada de", "ayud", "particip", "colabor",
    "apoy", "realic", "hice", "responsibilities included",
]

# Verbos de accion fuertes (un bullet que empieza asi es bueno)
VERBOS_FUERTES = {
    "led", "managed", "built", "created", "designed", "developed",
    "implemented", "reduced", "increased", "improved", "launched",
    "architected", "automated", "delivered", "drove", "spearheaded",
    "optimized", "achieved", "resolved", "detected", "mitigated",
    "analyzed", "coordinated", "standardized", "executed", "investigated",
    "monitored", "deployed", "configured", "engineered", "established",
    "streamlined", "accelerated", "eliminated", "negotiated", "mentored",
    "directed", "orchestrated", "transformed", "pioneered", "secured",
    "performed", "administered", "supported", "completed", "identified",
    "correlated", "represented", "reviewed", "hardened", "remediated",
    "lideré", "gestioné", "desarrollé", "implementé", "diseñé",
    "reduje", "aumenté", "mejoré", "construí", "automaticé", "coordiné",
    "detecté", "analicé", "resolví", "estandaricé",
}

# Verbos irregulares de accion (no terminan en -ed)
VERBOS_IRREGULARES = {
    "cut", "ran", "set", "met", "won", "kept", "sent", "made", "took",
    "gave", "held", "drove", "led", "built", "grew", "began", "rebuilt",
}


# ---------------------------------------------------------------------------
# Compatibilidad de texto / legibilidad ATS
# ---------------------------------------------------------------------------

# Fuentes seguras para ATS (segun Jobscan / Resume.io)
FUENTES_OK = {
    "arial", "calibri", "times new roman", "georgia", "cambria", "garamond",
    "helvetica", "palatino", "tahoma", "verdana", "book antiqua",
    "trebuchet ms", "century gothic", "lato", "segoe ui", "roboto",
    "open sans", "georgia pro", "constantia", "candara", "corbel",
}

# Caracteres tipograficos que algunos ATS leen mal
SMART_QUOTES = "“”‘’«»"          # " " ' ' « »
DASHES_RAROS = "–—"                                   # – —
# Bullets / simbolos decorativos (el • U+2022 se tolera; estos no)
SIMBOLOS_RAROS = ("▪●○■□‣⁃❖❦"
                  "➜➤★☆✔✘→»⁋❧")
EMOJI_RE = re.compile(
    "[\U0001F000-\U0001FAFF\U00002600-\U000027BF\U0001F1E6-\U0001F1FF⬀-⯿]")

_MESES_RE = (r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|"
             r"ene|abr|ago|dic|enero|febrero|marzo|abril|mayo|junio|julio|"
             r"agosto|septiembre|setiembre|octubre|noviembre|diciembre)")

# Acronimos tecnicos comunes -> forma completa (para sugerir deletrearlos)
ACRONIMOS_LARGOS = {
    "siem": "Security Information and Event Management",
    "edr": "Endpoint Detection and Response",
    "xdr": "Extended Detection and Response",
    "soc": "Security Operations Center",
    "soar": "Security Orchestration, Automation and Response",
    "iam": "Identity and Access Management",
    "pam": "Privileged Access Management",
    "dlp": "Data Loss Prevention",
    "ids": "Intrusion Detection System",
    "ips": "Intrusion Prevention System",
    "waf": "Web Application Firewall",
    "ioc": "Indicators of Compromise",
    "mfa": "Multi-Factor Authentication",
    "sso": "Single Sign-On",
    "vpn": "Virtual Private Network",
    "api": "Application Programming Interface",
    "sql": "Structured Query Language",
    "ci/cd": "Continuous Integration and Continuous Delivery",
    "osint": "Open-Source Intelligence",
    "ueba": "User and Entity Behavior Analytics",
    "grc": "Governance, Risk and Compliance",
    "rgpd": "Reglamento General de Protección de Datos",
    "ens": "Esquema Nacional de Seguridad",
    "lopd": "Ley Orgánica de Protección de Datos",
}

GENERICOS_ARCHIVO = {
    "resume", "cv", "curriculum", "curriculum vitae", "untitled", "document",
    "documento", "hoja de vida", "resumen", "final", "final version",
    "new resume", "my resume", "mi cv", "copy", "copia", "doc1", "sin titulo",
}


# ---------------------------------------------------------------------------
# Extraccion de texto (cuerpo, tablas, headers/footers)
# ---------------------------------------------------------------------------

def _texto_de_elemento(elem) -> List[str]:
    """Texto de todos los parrafos dentro de un elemento XML, en orden."""
    lineas = []
    for p_elem in elem.iter(qn("w:p")):
        t = "".join(r.text or "" for r in p_elem.iter(qn("w:t"))).strip()
        if t:
            lineas.append(t)
    return lineas


def _texto_cuerpo(doc) -> List[str]:
    """Texto del cuerpo del documento (parrafos + celdas de tabla) en orden."""
    lineas = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            t = "".join(r.text or "" for r in child.iter(qn("w:t"))).strip()
            if t:
                lineas.append(t)
        elif tag == "tbl":
            for row in child.iter(qn("w:tr")):
                for cell in row.iter(qn("w:tc")):
                    lineas.extend(_texto_de_elemento(cell))
    return lineas


def _texto_headers_footers(doc) -> List[str]:
    lineas = []
    for section in doc.sections:
        for cont in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            try:
                for p in cont.paragraphs:
                    if p.text.strip():
                        lineas.append(p.text.strip())
            except Exception:
                continue
    return lineas


# ---------------------------------------------------------------------------
# Helpers de check
# ---------------------------------------------------------------------------

def _check(estado: str, titulo: str, detalle: str) -> Dict:
    """estado: 'ok' | 'warning' | 'error'"""
    return {"estado": estado, "titulo": titulo, "detalle": detalle}


# ---------------------------------------------------------------------------
# Categoria: FORMATO (solo DOCX — requiere inspeccion del XML)
# ---------------------------------------------------------------------------

def _fuente_ok(f: str) -> bool:
    f = f.lower().replace(" light", "").replace(" bold", "").replace(" semibold", "").strip()
    return any(f == ok or f.startswith(ok) for ok in FUENTES_OK)


def _fuentes_usadas(doc) -> set:
    fuentes = set()
    try:
        n = doc.styles["Normal"].font.name
        if n:
            fuentes.add(n)
    except Exception:
        pass
    contenedores = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                contenedores.extend(cell.paragraphs)
    for p in contenedores:
        for r in p.runs:
            if r.font and r.font.name:
                fuentes.add(r.font.name)
    return fuentes


def _analizar_formato(doc) -> Tuple[List[Dict], int, int]:
    checks: List[Dict] = []
    max_pts = 30
    penalizacion = 0
    body = doc.element.body

    # Tablas
    n_tablas = sum(1 for _ in body.iter(qn("w:tbl")))
    if n_tablas:
        penalizacion += 10
        checks.append(_check(
            "warning", "Uso de tablas para el diseño",
            f"Tu CV usa {n_tablas} tabla(s). Muchos ATS (Taleo, iCIMS) leen las "
            "celdas en orden incorrecto o se saltan contenido. Si el layout es de "
            "columnas con tablas, considera una versión en una sola columna."))
    else:
        checks.append(_check(
            "ok", "Sin tablas", "El texto fluye en una sola columna, ideal para ATS."))

    # Imagenes
    n_img = sum(1 for _ in body.iter(qn("w:drawing"))) + \
            sum(1 for _ in body.iter(qn("w:pict")))
    if n_img:
        penalizacion += 15
        checks.append(_check(
            "error", "Imágenes o gráficos detectados",
            f"Se detectaron {n_img} elemento(s) gráfico(s). El texto dentro de "
            "imágenes es invisible para el ATS. Convierte cualquier dato (logos, "
            "íconos, gráficos de skills) a texto plano."))
    else:
        checks.append(_check(
            "ok", "Sin imágenes", "Todo el contenido es texto seleccionable."))

    # Text boxes
    n_txbx = sum(1 for _ in body.iter(qn("w:txbxContent")))
    if n_txbx:
        penalizacion += 8
        checks.append(_check(
            "warning", "Cuadros de texto (text boxes)",
            f"Se detectaron {n_txbx} cuadro(s) de texto. Su contenido suele "
            "perderse al ser parseado por el ATS. Pásalo al cuerpo del documento."))

    # Columnas multiples
    cols_multi = False
    for sectPr in body.iter(qn("w:sectPr")):
        for cols in sectPr.iter(qn("w:cols")):
            num = cols.get(qn("w:num"))
            if num and num.isdigit() and int(num) >= 2:
                cols_multi = True
    if cols_multi:
        penalizacion += 8
        checks.append(_check(
            "warning", "Diseño en múltiples columnas",
            "El documento usa columnas de sección. El ATS puede mezclar el orden "
            "de lectura. Una sola columna es lo más seguro."))

    # Contacto en header/footer
    hf = " ".join(_texto_headers_footers(doc))
    cuerpo = " ".join(_texto_cuerpo(doc))
    contacto_en_hf = bool(EMAIL_RE.search(hf) or PHONE_RE.search(hf))
    contacto_en_cuerpo = bool(EMAIL_RE.search(cuerpo) or PHONE_RE.search(cuerpo))
    if contacto_en_hf and not contacto_en_cuerpo:
        penalizacion += 12
        checks.append(_check(
            "error", "Datos de contacto en encabezado/pie",
            "Tu email o teléfono están en el header/footer. Muchos ATS ignoran "
            "esa zona — moverlos al cuerpo del CV es crítico."))
    elif contacto_en_hf:
        penalizacion += 4
        checks.append(_check(
            "warning", "Contacto también en encabezado/pie",
            "Tienes datos de contacto en el header/footer. Asegúrate de que "
            "también estén en el cuerpo (ya lo están), por si el ATS ignora esa zona."))

    # Fuentes no estandar
    fuentes = _fuentes_usadas(doc)
    no_ok = sorted({f for f in fuentes if not _fuente_ok(f)})
    if no_ok:
        penalizacion += 6
        checks.append(_check(
            "warning", "Fuentes no estándar",
            f"Detectamos: {', '.join(no_ok[:4])}. Usa fuentes seguras (Arial, "
            "Calibri, Times New Roman, Georgia, Garamond) — otras pueden parsearse "
            "como caracteres ilegibles."))
    else:
        checks.append(_check(
            "ok", "Fuentes estándar",
            "Las fuentes del documento son compatibles con ATS."))

    puntos = max(0, max_pts - penalizacion)
    return checks, puntos, max_pts


# ---------------------------------------------------------------------------
# Categoria: ESTRUCTURA
# ---------------------------------------------------------------------------

def _analizar_estructura(lineas: List[str]) -> Tuple[List[Dict], int, int]:
    checks: List[Dict] = []
    max_pts = 20
    # Solo las lineas cortas son candidatas a encabezado de seccion
    posibles_titulos = [l for l in lineas if len(l) < 60]

    faltantes = []
    for nombre, patron in SECCIONES.items():
        encontrada = any(patron.search(t) for t in posibles_titulos)
        if encontrada:
            checks.append(_check("ok", f"Sección «{nombre}»", "Presente y detectable."))
        else:
            faltantes.append(nombre)
            checks.append(_check(
                "warning", f"Falta la sección «{nombre}»",
                "El ATS busca encabezados estándar. Agrega esta sección con un "
                "título claro para que clasifique bien tu información."))

    penalizacion = len(faltantes) * 4
    puntos = max(0, max_pts - penalizacion)
    return checks, puntos, max_pts


# ---------------------------------------------------------------------------
# Categoria: CONTACTO
# ---------------------------------------------------------------------------

def _analizar_contacto(texto: str) -> Tuple[List[Dict], int, int]:
    checks: List[Dict] = []
    max_pts = 15
    penalizacion = 0

    if EMAIL_RE.search(texto):
        checks.append(_check("ok", "Email detectado", "Tu correo es legible para el ATS."))
    else:
        penalizacion += 7
        checks.append(_check(
            "error", "Sin email detectable",
            "No se encontró un email en formato estándar (nombre@dominio.com)."))

    if PHONE_RE.search(texto):
        checks.append(_check("ok", "Teléfono detectado", "Número de contacto presente."))
    else:
        penalizacion += 4
        checks.append(_check(
            "warning", "Sin teléfono detectable",
            "Agrega un número de teléfono con formato claro (+código país)."))

    if LINKEDIN_RE.search(texto):
        checks.append(_check("ok", "LinkedIn detectado", "Perfil profesional enlazado."))
    else:
        penalizacion += 4
        checks.append(_check(
            "warning", "Sin LinkedIn",
            "Incluye tu URL de LinkedIn (linkedin.com/in/tu-perfil)."))

    puntos = max(0, max_pts - penalizacion)
    return checks, puntos, max_pts


# ---------------------------------------------------------------------------
# Categoria: CONTENIDO
# ---------------------------------------------------------------------------

def _es_titulo_seccion(linea: str) -> bool:
    if len(linea) >= 60:
        return False
    return any(p.search(linea) for p in SECCIONES.values())


def _primera_palabra(linea: str) -> str:
    m = re.match(r"\s*([A-Za-zÁ-úÑñ]+)", linea)
    return m.group(1).lower() if m else ""


def _empieza_con_verbo_accion(linea: str) -> bool:
    """True si la linea empieza con un verbo de accion (fuerte o terminado en -ed)."""
    w = _primera_palabra(linea)
    if w in VERBOS_FUERTES or w in VERBOS_IRREGULARES:
        return True
    if len(w) > 3 and w.endswith("ed"):
        return True
    return False


def _empieza_con_verbo_debil(linea: str) -> bool:
    low = linea.lower().strip()
    return any(low.startswith(v) for v in VERBOS_DEBILES)


def _analizar_contenido(lineas: List[str]) -> Tuple[List[Dict], int, int]:
    checks: List[Dict] = []
    max_pts = 20
    penalizacion = 0

    # "Bullet de logro" = linea de longitud media que empieza con un verbo
    # (de accion o debil). Esto excluye el resumen, coursework, certificaciones.
    logros = [
        l for l in lineas
        if 30 <= len(l) <= 350 and not _es_titulo_seccion(l)
        and (_empieza_con_verbo_accion(l) or _empieza_con_verbo_debil(l))
    ]

    # 1. Verbos debiles
    debiles = [l for l in logros if _empieza_con_verbo_debil(l)]
    if debiles:
        prop = len(debiles) / max(1, len(logros))
        penalizacion += min(10, int(prop * 25) + 3)
        ejemplos = "; ".join(d[:50] + "…" for d in debiles[:2])
        checks.append(_check(
            "warning", f"{len(debiles)} frase(s) con verbos débiles",
            f"Evita «responsible for», «worked on», «helped». Usa verbos de "
            f"impacto (Led, Reduced, Built). Ej: {ejemplos}"))
    else:
        checks.append(_check(
            "ok", "Buen uso de verbos de acción",
            "Tus logros empiezan con verbos de impacto."))

    # 2. Metricas cuantificables
    con_metrica = [l for l in logros if re.search(r"\d+\s*%|\d[\d,.]*\s*\+?", l)]
    if logros:
        prop_con = len(con_metrica) / len(logros)
        if prop_con < 0.3:
            penalizacion += 10
            checks.append(_check(
                "warning", "Pocas métricas cuantificables",
                f"Solo {len(con_metrica)} de {len(logros)} logros tienen números. "
                "Agrega cifras: «Reduje el MTTR un 60%», «Analicé 100+ alertas/semana»."))
        elif prop_con < 0.6:
            penalizacion += 4
            checks.append(_check(
                "warning", "Métricas mejorables",
                f"{len(con_metrica)} de {len(logros)} logros tienen cifras. "
                "Cuantifica más resultados para destacar."))
        else:
            checks.append(_check(
                "ok", "Logros bien cuantificados",
                f"{len(con_metrica)} de {len(logros)} logros incluyen métricas."))

    # 3. Longitud total
    total_palabras = sum(len(l.split()) for l in lineas)
    if total_palabras < 200:
        penalizacion += 5
        checks.append(_check(
            "warning", "CV demasiado corto",
            f"~{total_palabras} palabras. Desarrolla cada experiencia con contexto, "
            "herramientas y resultados (apunta a 400–700)."))
    elif total_palabras > 1100:
        penalizacion += 3
        checks.append(_check(
            "warning", "CV extenso",
            f"~{total_palabras} palabras. Considera condensar a lo más relevante "
            "(idealmente 1–2 páginas)."))
    else:
        checks.append(_check(
            "ok", "Longitud adecuada",
            f"~{total_palabras} palabras, en el rango ideal."))

    puntos = max(0, max_pts - penalizacion)
    return checks, puntos, max_pts


# ---------------------------------------------------------------------------
# Categoria: LEGIBILIDAD ATS (fechas, caracteres, acronimos, nombre archivo)
# ---------------------------------------------------------------------------

def _analizar_legibilidad(texto: str, nombre_archivo: str) -> Tuple[List[Dict], int, int]:
    checks: List[Dict] = []
    max_pts = 15
    penalizacion = 0
    texto_low = texto.lower()

    # 1. Formato de fechas
    apostrofe = re.search(r"['’]\d{2}\b", texto)
    con_mes = len(re.findall(rf"{_MESES_RE}\.?\s+(?:19|20)\d{{2}}", texto, re.I))
    con_mes += len(re.findall(r"\b(?:0[1-9]|1[0-2])[/\-](?:19|20)\d{2}", texto))
    rangos_solo_anio = len(re.findall(
        r"\b(?:19|20)\d{2}\s*[-–—]\s*(?:(?:19|20)\d{2}|present|presente|actual|current)",
        texto, re.I))
    hay_fechas = con_mes or rangos_solo_anio or apostrofe
    if apostrofe:
        penalizacion += 5
        checks.append(_check(
            "warning", "Años abreviados con apóstrofe",
            "Detectamos fechas tipo «'21». El ATS puede no calcular bien tu "
            "experiencia. Usa el año completo: «2021» o «Ene 2021»."))
    elif rangos_solo_anio and con_mes == 0:
        penalizacion += 3
        checks.append(_check(
            "warning", "Fechas sin mes",
            "Tus fechas usan solo el año (2021–2023). Incluye el mes (Ene 2021 o "
            "01/2021) para que el ATS calcule con precisión tus años de experiencia."))
    elif hay_fechas:
        checks.append(_check(
            "ok", "Formato de fechas correcto",
            "Tus fechas son legibles y permiten calcular tu experiencia."))

    # 2. Caracteres especiales / simbolos
    simbolos = sorted({c for c in texto if c in SIMBOLOS_RAROS})
    emojis   = EMOJI_RE.findall(texto)
    leves    = any(c in texto for c in SMART_QUOTES) or any(c in texto for c in DASHES_RAROS)
    if simbolos or emojis:
        penalizacion += 5
        muestra = " ".join(simbolos[:5]) + (" + emojis" if emojis else "")
        checks.append(_check(
            "warning", "Símbolos o emojis no estándar",
            f"Detectamos: {muestra}. Algunos ATS los leen como «?» o los omiten. "
            "Usa viñetas estándar (• o -) y elimina íconos/emojis."))
    elif leves:
        penalizacion += 2
        checks.append(_check(
            "warning", "Comillas o guiones tipográficos",
            "Usas comillas curvas (“ ”) o guiones largos (– —). Los ATS antiguos "
            "pueden fallar; usa comillas rectas (\" ') y guion simple (-)."))
    else:
        checks.append(_check(
            "ok", "Sin caracteres problemáticos",
            "El texto usa caracteres estándar legibles por cualquier ATS."))

    # 3. Acronimos sin forma completa
    detectados, faltan = 0, []
    for acr, largo in ACRONIMOS_LARGOS.items():
        if re.search(rf"\b{re.escape(acr)}\b", texto_low):
            detectados += 1
            if largo.lower() not in texto_low:
                faltan.append(acr.upper())
    if faltan:
        penalizacion += min(4, len(faltan))
        ej = faltan[0]
        checks.append(_check(
            "warning", f"{len(faltan)} acrónimo(s) sin su forma completa",
            f"Escribe la forma larga al menos una vez: «{ACRONIMOS_LARGOS[ej.lower()]} "
            f"({ej})». ATS como Taleo no reconocen el acrónimo solo. "
            f"Faltan: {', '.join(faltan[:6])}."))
    elif detectados:
        checks.append(_check(
            "ok", "Acrónimos bien definidos",
            "Tus acrónimos técnicos incluyen su forma completa."))

    # 4. Nombre del archivo
    base = nombre_archivo.rsplit(".", 1)[0] if "." in nombre_archivo else nombre_archivo
    base = base.strip()
    probs = []
    if base.lower() in GENERICOS_ARCHIVO:
        probs.append("es genérico")
    if " " in base:
        probs.append("tiene espacios")
    if re.search(r"[^A-Za-z0-9_\- ]", base):
        probs.append("tiene acentos/caracteres especiales")
    if len(base) > 50:
        probs.append("es muy largo")
    if probs:
        penalizacion += 2
        checks.append(_check(
            "warning", "Nombre de archivo mejorable",
            f"El nombre «{base}» {', '.join(probs)}. Usa algo como "
            "«Nombre-Apellido-CV.docx» (solo letras y guiones)."))
    else:
        checks.append(_check(
            "ok", "Nombre de archivo correcto",
            "El nombre del archivo es claro y compatible."))

    puntos = max(0, max_pts - penalizacion)
    return checks, puntos, max_pts


# ---------------------------------------------------------------------------
# Nivel y resumen
# ---------------------------------------------------------------------------

def _nivel(score: int) -> Tuple[str, str]:
    if score >= 85:
        return "Excelente", "verde"
    if score >= 70:
        return "Bueno", "verde"
    if score >= 50:
        return "Regular", "amarillo"
    return "Necesita trabajo", "rojo"


# ---------------------------------------------------------------------------
# Funcion publica
# ---------------------------------------------------------------------------

def analizar_ats(contenido: bytes, nombre_archivo: str) -> Dict:
    """
    Analiza un CV (.docx o .pdf) y devuelve un reporte de compatibilidad ATS.
    Para PDF, el análisis de formato profundo no aplica (no hay XML inspeccionable).
    """
    ext = (nombre_archivo or "").rsplit(".", 1)[-1].lower()
    categorias: List[Dict] = []

    if ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(contenido))
        lineas = _texto_cuerpo(doc)
        texto  = "\n".join(lineas + _texto_headers_footers(doc))

        c_fmt, p_fmt, m_fmt = _analizar_formato(doc)
        c_est, p_est, m_est = _analizar_estructura(lineas)
        c_con, p_con, m_con = _analizar_contacto(texto)
        c_cnt, p_cnt, m_cnt = _analizar_contenido(lineas)
        c_leg, p_leg, m_leg = _analizar_legibilidad(texto, nombre_archivo)

        categorias = [
            {"nombre": "Formato",    "icono": "📐", "puntos": p_fmt, "max_puntos": m_fmt, "checks": c_fmt},
            {"nombre": "Estructura", "icono": "🗂️", "puntos": p_est, "max_puntos": m_est, "checks": c_est},
            {"nombre": "Contacto",   "icono": "📇", "puntos": p_con, "max_puntos": m_con, "checks": c_con},
            {"nombre": "Contenido",  "icono": "✍️", "puntos": p_cnt, "max_puntos": m_cnt, "checks": c_cnt},
            {"nombre": "Legibilidad ATS", "icono": "🔤", "puntos": p_leg, "max_puntos": m_leg, "checks": c_leg},
        ]
        tipo = "docx"

    elif ext == "pdf":
        from app.services.extractor import _desde_pdf
        texto_pdf = _desde_pdf(contenido)
        lineas = [l.strip() for l in texto_pdf.splitlines() if l.strip()]
        texto  = texto_pdf

        c_est, p_est, m_est = _analizar_estructura(lineas)
        c_con, p_con, m_con = _analizar_contacto(texto)
        c_cnt, p_cnt, m_cnt = _analizar_contenido(lineas)
        c_leg, p_leg, m_leg = _analizar_legibilidad(texto, nombre_archivo)

        nota_fmt = [_check(
            "warning", "Análisis de formato limitado en PDF",
            "Para revisar tablas, columnas, imágenes, fuentes y headers/footers, sube "
            "tu CV en formato DOCX. En PDF analizamos estructura, contacto, contenido y legibilidad.")]
        categorias = [
            {"nombre": "Formato",    "icono": "📐", "puntos": 0, "max_puntos": 0, "checks": nota_fmt},
            {"nombre": "Estructura", "icono": "🗂️", "puntos": p_est, "max_puntos": m_est, "checks": c_est},
            {"nombre": "Contacto",   "icono": "📇", "puntos": p_con, "max_puntos": m_con, "checks": c_con},
            {"nombre": "Contenido",  "icono": "✍️", "puntos": p_cnt, "max_puntos": m_cnt, "checks": c_cnt},
            {"nombre": "Legibilidad ATS", "icono": "🔤", "puntos": p_leg, "max_puntos": m_leg, "checks": c_leg},
        ]
        tipo = "pdf"

    else:
        from fastapi import HTTPException
        raise HTTPException(
            status_code=415,
            detail=f"Formato no soportado: .{ext}. Sube un archivo .docx o .pdf")

    return _ensamblar_reporte(categorias, tipo)


def _ensamblar_reporte(categorias: List[Dict], tipo: str) -> Dict:
    total_pts = sum(c["puntos"] for c in categorias)
    total_max = sum(c["max_puntos"] for c in categorias)
    score = int(round(total_pts / total_max * 100)) if total_max else 0
    nivel, color = _nivel(score)

    n_errores = sum(1 for c in categorias for ch in c["checks"] if ch["estado"] == "error")
    n_warn    = sum(1 for c in categorias for ch in c["checks"] if ch["estado"] == "warning")

    if n_errores == 0 and n_warn == 0:
        resumen = "Tu CV es muy compatible con ATS. ¡Excelente trabajo!"
    elif n_errores:
        resumen = (f"Se detectaron {n_errores} problema(s) crítico(s) y {n_warn} "
                   "advertencia(s). Corrige primero los marcados en rojo.")
    else:
        resumen = (f"Sin problemas críticos. Hay {n_warn} mejora(s) sugerida(s) "
                   "para subir tu compatibilidad ATS.")

    return {
        "score": score,
        "nivel": nivel,
        "color": color,
        "resumen": resumen,
        "tipo_archivo": tipo,
        "n_errores": n_errores,
        "n_advertencias": n_warn,
        "categorias": categorias,
    }


def analizar_ats_texto(texto: str) -> Dict:
    """
    Analiza un CV pegado como TEXTO PLANO. El análisis de formato profundo (tablas,
    columnas, fuentes) no aplica sin el archivo, pero sí estructura, contacto,
    contenido y legibilidad.
    """
    lineas = [l.strip() for l in texto.splitlines() if l.strip()]
    c_est, p_est, m_est = _analizar_estructura(lineas)
    c_con, p_con, m_con = _analizar_contacto(texto)
    c_cnt, p_cnt, m_cnt = _analizar_contenido(lineas)
    c_leg, p_leg, m_leg = _analizar_legibilidad(texto, "cv.txt")
    nota_fmt = [_check(
        "warning", "Análisis de formato limitado (texto plano)",
        "Pegaste el CV como texto, así que no revisamos tablas, columnas, imágenes "
        "ni fuentes. Para el análisis de formato completo, sube el DOCX.")]
    categorias = [
        {"nombre": "Formato",    "icono": "📐", "puntos": 0, "max_puntos": 0, "checks": nota_fmt},
        {"nombre": "Estructura", "icono": "🗂️", "puntos": p_est, "max_puntos": m_est, "checks": c_est},
        {"nombre": "Contacto",   "icono": "📇", "puntos": p_con, "max_puntos": m_con, "checks": c_con},
        {"nombre": "Contenido",  "icono": "✍️", "puntos": p_cnt, "max_puntos": m_cnt, "checks": c_cnt},
        {"nombre": "Legibilidad ATS", "icono": "🔤", "puntos": p_leg, "max_puntos": m_leg, "checks": c_leg},
    ]
    return _ensamblar_reporte(categorias, "texto")
