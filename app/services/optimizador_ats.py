"""
Optimizador ATS — pipeline completo para score 95-100%.

    EXTRAER -> MAPEAR (JSON) -> ENRIQUECER (JD opcional) -> RENDERIZAR -> VERIFICAR

A diferencia de la plantilla simple, el renderizador construye el documento
100% desde el JSON intermedio (nunca desde las lineas del original), por lo que
es imposible heredar tablas, columnas, iconos, fuentes raras o fechas ambiguas.

Reglas estrictas aplicadas (segun practicas documentadas de Taleo, Workday,
SuccessFactors y Greenhouse):
  - Una columna, fuente estandar, vinetas simples, contacto en el cuerpo.
  - Secciones renombradas al canon universal (ES/EN segun el idioma del CV).
  - Fechas normalizadas a "Mes AAAA" y experiencia en orden cronologico inverso.
  - Con vacante: inyeccion de titulo del puesto, keywords, acronimos expandidos
    y metricas de impacto en bullets (placeholders ~ ajustables).
"""

import io
import re
from datetime import datetime
from typing import Dict, List, Tuple

from app.services.adaptador import (
    _keywords_de, _analizar_cobertura, _kw_cubierta, norm_alias,
    _detectar_titulo_vacante, _titulo_en_cv,
)
from app.services.adaptador_docx import _construir_resumen_completo, _formato_keyword
from app.services.ats_checker import ACRONIMOS_LARGOS
from app.services.mejorador_bullets import (
    _es_bullet_debil, _reescribir_bullet, tiene_metrica, sufijo_metrica, detectar_sector,
)
from app.services.parser_ats import (
    EMAIL_RE, PHONE_RE, LINKEDIN_RE, UBIC_RE, PRESENTE, _detectar_nombre,
)
from app.services.plantilla_ats import (
    _CANON, _idioma, _limpiar, _es_header_seccion, _es_linea_contacto_simple,
    _VERBO_BULLET, _parece_subtitulo,
    _add_section_header, _add_paragraph, _add_subtitulo, _add_bullet,
)


# ---------------------------------------------------------------------------
# Normalizacion de fechas -> "Mes AAAA"
# ---------------------------------------------------------------------------

_MES_NUM = {"ene": 1, "jan": 1, "feb": 2, "mar": 3, "apr": 4, "abr": 4,
            "may": 5, "jun": 6, "jul": 7, "aug": 8, "ago": 8, "sep": 9,
            "set": 9, "oct": 10, "nov": 11, "dec": 12, "dic": 12}

_ABBR = {
    "es": ["Ene", "Feb", "Mar", "Abr", "May", "Jun",
           "Jul", "Ago", "Sep", "Oct", "Nov", "Dic"],
    "en": ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
           "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"],
}
_PRES_TXT = {"es": "Presente", "en": "Present"}

_MES_TXT = (r"(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|"
            r"setiembre|octubre|noviembre|diciembre|january|february|march|"
            r"april|june|july|august|september|october|november|december|"
            r"jan|ene|feb|mar|abr|apr|may|jun|jul|ago|aug|sept|sep|set|"
            r"oct|nov|dic|dec)")

# Token de fecha: "Mes AAAA" | "MM/AAAA" | "'AA" | "AAAA" | presente
_TOK_TXT = (rf"(?:{_MES_TXT}\.?\s*(?:de\s+)?(?:19|20)\d{{2}}"
            rf"|(?:0?[1-9]|1[0-2])[/](?:19|20)\d{{2}}"
            rf"|'\d{{2}}"
            rf"|(?:19|20)\d{{2}}"
            rf"|{PRESENTE})")

_RANGO_X = re.compile(
    rf"({_TOK_TXT})(?:\s*[-–—]\s*|\s+(?:to|hasta|al|a)\s+)({_TOK_TXT})",
    re.IGNORECASE)


def _parse_token(tok: str) -> Tuple[str, int, int]:
    """Devuelve (tipo, anio, mes). tipo: 'pres' | 'fecha' | ''."""
    t = tok.strip().lower().rstrip(".")
    if re.fullmatch(PRESENTE, t, re.IGNORECASE):
        return ("pres", 0, 0)
    anio = 0
    m = re.search(r"(?:19|20)\d{2}", t)
    if m:
        anio = int(m.group(0))
    else:
        m2 = re.fullmatch(r"'(\d{2})", t)
        if m2:
            yy = int(m2.group(1))
            anio = 2000 + yy if yy <= (datetime.now().year - 2000 + 1) else 1900 + yy
    mes = 0
    m3 = re.match(r"([a-záéíóú]+)", t)
    if m3 and m3.group(1)[:3] in _MES_NUM:
        mes = _MES_NUM[m3.group(1)[:3]]
    if not mes:
        m4 = re.match(r"(0?[1-9]|1[0-2])/", t)
        if m4:
            mes = int(m4.group(1))
    return ("fecha" if anio else "", anio, mes)


def _fmt_token(tipo: str, anio: int, mes: int, idioma: str) -> str:
    if tipo == "pres":
        return _PRES_TXT[idioma]
    if anio and mes:
        return f"{_ABBR[idioma][mes - 1]} {anio}"
    if anio:
        return str(anio)
    return ""


def _normalizar_fechas(texto: str, idioma: str) -> Tuple[str, int]:
    """Normaliza todos los rangos de fecha de un texto. Devuelve (texto, n_cambios)."""
    cambios = 0

    def _rep(m):
        nonlocal cambios
        a = _fmt_token(*_parse_token(m.group(1)), idioma)
        b = _fmt_token(*_parse_token(m.group(2)), idioma)
        if not a or not b:
            return m.group(0)
        nuevo = f"{a} - {b}"
        if nuevo != m.group(0):
            cambios += 1
        return nuevo

    return _RANGO_X.sub(_rep, texto), cambios


def _clave_fin(periodo: str) -> Tuple[int, int]:
    """Clave de orden (anio_fin, mes_fin) para orden cronologico inverso."""
    toks = re.findall(_TOK_TXT, periodo or "", flags=re.IGNORECASE)
    if not toks:
        return (-1, 0)
    tipo, anio, mes = _parse_token(toks[-1])
    if tipo == "pres":
        return (9999, 12)
    return (anio or -1, mes)


# ---------------------------------------------------------------------------
# Etapa 2 — MAPEAR: texto -> JSON intermedio
# ---------------------------------------------------------------------------

_LISTAS = ("educacion", "habilidades", "certificaciones", "proyectos",
           "logros", "idiomas", "referencias")


_LABEL_CONTACTO = re.compile(
    r"^(location|ubicaci[oó]n|phone|tel[eé]fono|email|correo|address|"
    r"direcci[oó]n|linkedin)\b", re.IGNORECASE)


def _es_contacto_en_cuerpo(linea: str) -> bool:
    """
    Linea de contacto dentro de una seccion (para no duplicarla).
    NUNCA considera contacto una linea con rango de fechas — PHONE_RE da
    falsos positivos con rangos de años tipo «2017 - 2019».
    """
    if len(linea) > 60 or _RANGO_X.search(linea):
        return False
    return bool(EMAIL_RE.search(linea) or LINKEDIN_RE.search(linea)
                or _LABEL_CONTACTO.match(linea))


def _parece_empresa(linea: str) -> bool:
    t = linea.strip()
    return (len(t) <= 120 and not t.endswith(".")
            and not _VERBO_BULLET.match(t)
            and not _es_bullet_debil(t))


def _parsear_puestos(lineas: List[str], idioma: str) -> Tuple[List[Dict], int]:
    """Convierte las lineas de la seccion experiencia en puestos estructurados."""
    puestos: List[Dict] = []
    actual = None
    n_fechas = 0
    for linea in lineas:
        m = _RANGO_X.search(linea)
        if m:
            titulo = linea[:m.start()].strip(" \t|-–—:·")
            periodo, nf = _normalizar_fechas(m.group(0), idioma)
            n_fechas += nf
            actual = {"titulo": titulo or _limpiar(linea), "empresa": "",
                      "periodo": periodo, "bullets": []}
            puestos.append(actual)
            continue
        if actual is None:
            actual = {"titulo": _limpiar(linea), "empresa": "", "periodo": "",
                      "bullets": []}
            puestos.append(actual)
            continue
        if not actual["empresa"] and not actual["bullets"] and _parece_empresa(linea):
            actual["empresa"] = _limpiar(linea)
        else:
            actual["bullets"].append(_limpiar(linea))
    return puestos, n_fechas


def mapear_cv(texto_cv: str) -> Tuple[Dict, int]:
    """Mapea el CV a un objeto estructurado. Devuelve (mapa, fechas_normalizadas)."""
    idioma = _idioma(texto_cv)
    lineas = [l.strip() for l in texto_cv.splitlines() if l.strip()]

    nombre = _detectar_nombre(lineas) or (lineas[0] if lineas else "")
    email = EMAIL_RE.search(texto_cv)
    tel = PHONE_RE.search(texto_cv)
    linked = LINKEDIN_RE.search(texto_cv)
    ubic = UBIC_RE.search(texto_cv)

    mapa: Dict = {
        "idioma": idioma,
        "nombre": _limpiar(nombre),
        "headline": "",
        "contacto": {
            "email": email.group(0) if email else "",
            "telefono": tel.group(0).strip() if tel else "",
            "linkedin": linked.group(0) if linked else "",
            "ubicacion": ubic.group(1).strip()[:60] if ubic else "",
        },
        "resumen": "",
        "experiencia": [],
        "otros": [],
    }
    for k in _LISTAS:
        mapa[k] = []

    # Segmentar por encabezados reales conocidos
    secciones: List[Dict] = []
    actual = None
    nombre_low = mapa["nombre"].lower()
    for linea in lineas:
        tipo = _es_header_seccion(linea)
        if tipo:
            actual = {"tipo": tipo, "header": linea, "lineas": []}
            secciones.append(actual)
            continue
        if actual is None:
            if (not mapa["headline"] and linea.lower() != nombre_low
                    and not _es_linea_contacto_simple(linea) and len(linea) <= 90):
                mapa["headline"] = _limpiar(linea)
            continue
        actual["lineas"].append(linea)

    n_fechas_total = 0
    for sec in secciones:
        tipo = sec["tipo"]
        cuerpo = [l for l in sec["lineas"]
                  if l.lower() != nombre_low
                  and not (tipo != "contacto" and _es_contacto_en_cuerpo(l))]
        if tipo == "contacto" or not cuerpo:
            continue
        if tipo == "resumen":
            mapa["resumen"] = (mapa["resumen"] + " " + " ".join(cuerpo)).strip()
        elif tipo == "experiencia":
            puestos, nf = _parsear_puestos(cuerpo, idioma)
            mapa["experiencia"].extend(puestos)
            n_fechas_total += nf
        elif tipo in _LISTAS:
            for l in cuerpo:
                ln, nf = _normalizar_fechas(_limpiar(l), idioma)
                n_fechas_total += nf
                mapa[tipo].append(ln)
        else:
            mapa["otros"].append({"header": sec["header"], "lineas": cuerpo})

    return mapa, n_fechas_total


def _texto_de_mapa(m: Dict) -> str:
    partes = [m["nombre"], m["headline"], m["resumen"]]
    for j in m["experiencia"]:
        partes += [j["titulo"], j["empresa"], j["periodo"]] + j["bullets"]
    for k in _LISTAS:
        partes += m[k]
    for o in m["otros"]:
        partes += [o["header"]] + o["lineas"]
    partes += [v for v in m["contacto"].values() if v]
    return "\n".join(p for p in partes if p)


# ---------------------------------------------------------------------------
# Etapa 3 — ENRIQUECER: verbos, metricas, keywords de la vacante, acronimos
# ---------------------------------------------------------------------------

def _expandir_acronimos_mapa(mapa: Dict) -> int:
    texto = _texto_de_mapa(mapa).lower()
    n = 0
    for acr, largo in ACRONIMOS_LARGOS.items():
        if largo.lower() in texto:
            continue
        patron = re.compile(rf"\b{re.escape(acr)}\b", re.IGNORECASE)
        if not patron.search(texto):
            continue

        def _rep(m):
            return f"{largo} ({m.group(0).upper()})"

        hecho = False
        if patron.search(mapa["resumen"]):
            mapa["resumen"] = patron.sub(_rep, mapa["resumen"], count=1)
            hecho = True
        if not hecho:
            for i, h in enumerate(mapa["habilidades"]):
                if patron.search(h):
                    mapa["habilidades"][i] = patron.sub(_rep, h, count=1)
                    hecho = True
                    break
        if not hecho:
            for job in mapa["experiencia"]:
                for i, b in enumerate(job["bullets"]):
                    if patron.search(b):
                        job["bullets"][i] = patron.sub(_rep, b, count=1)
                        hecho = True
                        break
                if hecho:
                    break
        if hecho:
            n += 1
    return n


def enriquecer(mapa: Dict, vacante_texto: str) -> List[Dict]:
    """Aplica las mejoras de contenido sobre el JSON. Devuelve cambios tipados y
    además guarda un resumen estructurado en mapa['_meta'] (keywords inyectadas…)."""
    cambios: List[Dict] = []
    meta = {"keywords_inyectadas": [], "titulo_inyectado": None,
            "n_verbos": 0, "n_metricas": 0, "n_acronimos": 0}

    # 1. Bullets de experiencia: verbos de accion + metricas de impacto
    n_verb = n_met = 0
    sector = detectar_sector(_texto_de_mapa(mapa))
    contadores: Dict[str, int] = {}
    for job in mapa["experiencia"]:
        nuevos = []
        for b in job["bullets"]:
            if _es_bullet_debil(b):
                rb = _reescribir_bullet(b)
                if rb and rb.strip() != b.strip():
                    b = rb
                    n_verb += 1
            if len(b) >= 40 and not tiene_metrica(b):
                b = b.rstrip().rstrip(".") + ", " + sufijo_metrica(b, contadores, sector) + "."
                n_met += 1
            nuevos.append(b)
        job["bullets"] = nuevos
    meta["n_verbos"], meta["n_metricas"] = n_verb, n_met
    if n_verb:
        cambios.append(_cambio("verbos", f"{n_verb} bullet(s) reescritos con verbos de acción"))
    if n_met:
        cambios.append(_cambio("metricas",
                               f"{n_met} métrica(s) de impacto añadidas — ajusta los valores con ~"))

    # 2. Vacante: titulo + keywords en resumen y habilidades
    if vacante_texto.strip():
        texto_total = _texto_de_mapa(mapa)
        kw_vacante = _keywords_de(vacante_texto)[:30]
        cubiertas, sugeridas = _analizar_cobertura(kw_vacante, texto_total)

        titulo = _detectar_titulo_vacante(vacante_texto)
        if titulo and not _titulo_en_cv(titulo, texto_total):
            mapa["headline"] = (mapa["headline"] + " | " + titulo).strip(" |")
            meta["titulo_inyectado"] = titulo
            cambios.append(_cambio("cargo", f"Cargo objetivo «{titulo}» añadido al titular"))

        base = mapa["resumen"] or mapa["headline"]
        mapa["resumen"] = _construir_resumen_completo(base, cubiertas, sugeridas)

        cv_alias = norm_alias(_texto_de_mapa(mapa))
        faltantes = [k for k in (cubiertas + sugeridas)
                     if not _kw_cubierta(cv_alias, k)]
        if faltantes:
            mapa["habilidades"].append(
                " | ".join(_formato_keyword(k) for k in faltantes))
        # Keywords concretas de la vacante que se inyectaron (sugeridas + faltantes).
        inyectadas: List[str] = []
        for k in (sugeridas + faltantes):
            kf = _formato_keyword(k)
            if kf and kf not in inyectadas:
                inyectadas.append(kf)
        meta["keywords_inyectadas"] = inyectadas
        if sugeridas or faltantes:
            cambios.append(_cambio("keywords",
                f"{len(inyectadas)} keyword(s) de la vacante integradas en resumen y habilidades"))

    # 3. Acronimos -> forma completa (regla Taleo)
    n_acr = _expandir_acronimos_mapa(mapa)
    meta["n_acronimos"] = n_acr
    if n_acr:
        cambios.append(_cambio("acronimos", f"{n_acr} acrónimo(s) expandidos a su forma completa"))

    mapa["_meta"] = meta
    return cambios


# ---------------------------------------------------------------------------
# Etapa 4 — RENDERIZAR: JSON -> DOCX perfecto
# ---------------------------------------------------------------------------

_ORDEN_SECCIONES = ["resumen", "experiencia", "educacion", "habilidades",
                    "certificaciones", "proyectos", "logros", "idiomas",
                    "referencias"]


def _add_lineas_render(doc, lineas: List[str]):
    for t in lineas:
        if not t:
            continue
        if _parece_subtitulo(t):
            _add_subtitulo(doc, t)
        else:
            _add_bullet(doc, t)


def renderizar_docx(mapa: Dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    idioma = mapa["idioma"]
    canon = _CANON[idioma]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Nombre + headline
    if mapa["nombre"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(mapa["nombre"].upper())
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0x12, 0x2A, 0x4A)
    if mapa["headline"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(mapa["headline"])
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Contacto con encabezado estandar (los ATS mapean la seccion por regex)
    contacto = [v for v in (mapa["contacto"]["email"], mapa["contacto"]["telefono"],
                            mapa["contacto"]["linkedin"], mapa["contacto"]["ubicacion"]) if v]
    if contacto:
        _add_section_header(doc, canon["contacto"])
        _add_paragraph(doc, "  |  ".join(contacto))

    # Experiencia en orden cronologico inverso
    experiencia = sorted(
        enumerate(mapa["experiencia"]),
        key=lambda par: (_clave_fin(par[1]["periodo"]), -par[0]),
        reverse=True)
    experiencia = [j for _, j in experiencia]

    for tipo in _ORDEN_SECCIONES:
        if tipo == "resumen":
            if mapa["resumen"]:
                _add_section_header(doc, canon["resumen"])
                _add_paragraph(doc, mapa["resumen"])
        elif tipo == "experiencia":
            if experiencia:
                _add_section_header(doc, canon["experiencia"])
                for job in experiencia:
                    _add_subtitulo(doc, job["titulo"])
                    linea2 = "  |  ".join(x for x in (job["empresa"], job["periodo"]) if x)
                    if linea2:
                        _add_paragraph(doc, linea2)
                    for b in job["bullets"]:
                        _add_bullet(doc, b)
        elif mapa[tipo]:
            _add_section_header(doc, canon[tipo])
            if tipo == "habilidades":
                for l in mapa[tipo]:
                    _add_bullet(doc, l)
            else:
                _add_lineas_render(doc, mapa[tipo])

    for otro in mapa["otros"]:
        _add_section_header(doc, _limpiar(otro["header"]).upper().rstrip(":"))
        _add_lineas_render(doc, otro["lineas"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def renderizar_texto(mapa: Dict) -> str:
    """Misma reconstrucción que el DOCX pero en texto plano, para la vista previa."""
    canon = _CANON[mapa["idioma"]]
    out: List[str] = []
    if mapa["nombre"]:
        out.append(mapa["nombre"].upper())
    if mapa["headline"]:
        out.append(mapa["headline"])

    contacto = [v for v in (mapa["contacto"]["email"], mapa["contacto"]["telefono"],
                            mapa["contacto"]["linkedin"], mapa["contacto"]["ubicacion"]) if v]
    if contacto:
        out += ["", canon["contacto"], "  |  ".join(contacto)]

    experiencia = [j for _, j in sorted(
        enumerate(mapa["experiencia"]),
        key=lambda par: (_clave_fin(par[1]["periodo"]), -par[0]), reverse=True)]

    for tipo in _ORDEN_SECCIONES:
        if tipo == "resumen":
            if mapa["resumen"]:
                out += ["", canon["resumen"], mapa["resumen"]]
        elif tipo == "experiencia":
            if experiencia:
                out += ["", canon["experiencia"]]
                for job in experiencia:
                    out.append(job["titulo"])
                    linea2 = "  |  ".join(x for x in (job["empresa"], job["periodo"]) if x)
                    if linea2:
                        out.append(linea2)
                    out += ["• " + b for b in job["bullets"]]
        elif mapa[tipo]:
            out += ["", canon[tipo]]
            out += [("• " + l) if tipo == "habilidades" else l for l in mapa[tipo]]

    for otro in mapa["otros"]:
        out += ["", _limpiar(otro["header"]).upper().rstrip(":")]
        out += otro["lineas"]
    return "\n".join(out).strip()


def _cambio(tipo: str, texto: str) -> Dict:
    """Cambio tipado para colorear/legendar en la vista previa."""
    return {"tipo": tipo, "texto": texto}


# ---------------------------------------------------------------------------
# Funcion publica — pipeline completo
# ---------------------------------------------------------------------------

def optimizar_cv(texto_cv: str, vacante_texto: str = "") -> Dict:
    """
    Ejecuta el pipeline Mapear -> Enriquecer -> Renderizar.
    Devuelve {"docx": bytes, "mapa": dict, "cambios": [str]}.
    """
    mapa, n_fechas = mapear_cv(texto_cv)

    cambios = [_cambio("estructura", "Estructura reconstruida: una columna, sin tablas, gráficos ni iconos"),
               _cambio("secciones", "Secciones renombradas al estándar universal ATS")]
    if n_fechas:
        cambios.append(_cambio("fechas", f"{n_fechas} fecha(s) normalizadas al formato Mes AAAA"))

    orden_original = [j["titulo"] for j in mapa["experiencia"]]
    cambios += enriquecer(mapa, vacante_texto)

    docx = renderizar_docx(mapa)
    texto_optimizado = renderizar_texto(mapa)

    orden_final = [j["titulo"] for _, j in sorted(
        enumerate(mapa["experiencia"]),
        key=lambda par: (_clave_fin(par[1]["periodo"]), -par[0]),
        reverse=True)]
    if len(orden_original) > 1 and orden_original != orden_final:
        cambios.insert(2, _cambio("orden", "Experiencia reordenada cronológicamente (más reciente primero)"))

    return {
        "docx": docx,
        "mapa": mapa,
        "cambios": cambios,
        "texto_optimizado": texto_optimizado,
        "keywords_inyectadas": mapa.get("_meta", {}).get("keywords_inyectadas", []),
    }
